"""DOCX to Markdown converter for doc2md."""

import os
import re
import sys
from datetime import datetime
from pathlib import Path

from utils import VERSION
from utils.images import _check_content_sufficient
from utils.tables import _format_cell_value


def _extract_docx_images(doc, input_path, output_dir, config):
    """Extract embedded images from a .docx document.
    Returns {rId: "![](./images/imgN.ext)"} mapping."""
    import zipfile

    if not config.get("images", {}).get("extract", True):
        return {}

    img_subfolder = config.get("images", {}).get("subfolder", "images")
    min_w = config.get("images", {}).get("min_width", 100)
    min_h = config.get("images", {}).get("min_height", 100)
    img_dir = output_dir / img_subfolder
    img_dir.mkdir(parents=True, exist_ok=True)

    # EMU to pixel conversion (1 inch = 914400 EMU, 96 DPI)
    EMU_PER_PIXEL = 914400 / 96

    rid_to_md = {}
    img_idx = 0

    # Build rId -> image part mapping from relationships
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                blob = rel.target_part.blob
                ext = os.path.splitext(rel.target_part.partname)[-1].lstrip('.')
                if not ext:
                    ext = 'png'
            except Exception:
                continue

            img_idx += 1
            img_filename = f"img{img_idx}.{ext}"
            img_path = img_dir / img_filename
            with open(img_path, 'wb') as f:
                f.write(blob)

            rid_to_md[rel_id] = f"![image](./{img_subfolder}/{img_filename})"

    # Now check sizes from drawing XML and filter small images
    # Parse document XML to find <wp:extent> for each image
    from xml.etree import ElementTree as ET
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }

    rid_sizes = {}
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            # Find all drawing/blip references in this paragraph
            for drawing in element.iter(f'{{{nsmap["wp"]}}}extent'):
                cx = int(drawing.get('cx', '0'))
                cy = int(drawing.get('cy', '0'))
                # Find associated blip rId
                parent = element
                for blip in parent.iter(f'{{{nsmap["a"]}}}blip'):
                    rid = blip.get(f'{{{nsmap["r"]}}}embed')
                    if rid:
                        rid_sizes[rid] = (cx / EMU_PER_PIXEL, cy / EMU_PER_PIXEL)

    # Filter out small images and remove their files
    for rid, (w, h) in rid_sizes.items():
        if rid in rid_to_md and (w < min_w or h < min_h):
            # Extract filename from markdown ref to delete the file
            md_ref = rid_to_md[rid]
            fname = md_ref.split('/')[-1].rstrip(')')
            img_path = img_dir / fname
            if img_path.exists():
                img_path.unlink()
            del rid_to_md[rid]

    return rid_to_md


def _extract_chart_data(input_path):
    """Extract chart data from .docx as markdown tables.
    Returns {chart_filename: markdown_table_string}."""
    import zipfile
    from xml.etree import ElementTree as ET

    C_NS = 'http://schemas.openxmlformats.org/drawingml/2006/chart'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    chart_tables = {}

    try:
        with zipfile.ZipFile(str(input_path), 'r') as zf:
            chart_files = [n for n in zf.namelist() if n.startswith('word/charts/chart') and n.endswith('.xml')]

            for chart_file in chart_files:
                chart_name = os.path.basename(chart_file)
                try:
                    chart_xml = zf.read(chart_file)
                    root = ET.fromstring(chart_xml)

                    # Extract series data
                    series_list = []
                    categories = []

                    for ser in root.iter(f'{{{C_NS}}}ser'):
                        # Series name
                        ser_name = ''
                        tx = ser.find(f'{{{C_NS}}}tx')
                        if tx is not None:
                            sv = tx.find(f'.//{{{C_NS}}}v')
                            if sv is not None and sv.text:
                                ser_name = sv.text

                        # Categories (from first series only)
                        if not categories:
                            cat = ser.find(f'{{{C_NS}}}cat')
                            if cat is not None:
                                for v in cat.iter(f'{{{C_NS}}}v'):
                                    if v.text:
                                        categories.append(v.text)

                        # Values
                        values = []
                        val = ser.find(f'{{{C_NS}}}val')
                        if val is not None:
                            for v in val.iter(f'{{{C_NS}}}v'):
                                values.append(v.text if v.text else '')

                        if ser_name or values:
                            series_list.append((ser_name or f'Series {len(series_list)+1}', values))

                    if not series_list:
                        continue

                    # Build markdown table
                    if categories:
                        header = ['Category'] + [s[0] for s in series_list]
                        rows = []
                        for i, cat in enumerate(categories):
                            row = [cat]
                            for _, vals in series_list:
                                row.append(vals[i] if i < len(vals) else '')
                            rows.append(row)
                    else:
                        header = [s[0] for s in series_list]
                        max_len = max(len(s[1]) for s in series_list) if series_list else 0
                        rows = []
                        for i in range(max_len):
                            row = []
                            for _, vals in series_list:
                                row.append(vals[i] if i < len(vals) else '')
                            rows.append(row)

                    if rows:
                        max_cols = len(header)
                        table_lines = ['| ' + ' | '.join(header) + ' |']
                        table_lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
                        for row in rows:
                            while len(row) < max_cols:
                                row.append('')
                            table_lines.append('| ' + ' | '.join(row[:max_cols]) + ' |')
                        chart_tables[chart_name] = '\n'.join(table_lines)

                except Exception:
                    continue

            # Also try embedded xlsx in charts
            for chart_file in chart_files:
                chart_name = os.path.basename(chart_file)
                if chart_name in chart_tables:
                    continue  # Already got data from XML
                # Check for embedded xlsx
                for embed_name in zf.namelist():
                    if embed_name.startswith('word/embeddings/') and embed_name.endswith('.xlsx'):
                        try:
                            import openpyxl
                            import io
                            xlsx_data = zf.read(embed_name)
                            wb = openpyxl.load_workbook(io.BytesIO(xlsx_data), data_only=True)
                            ws = wb.active
                            if ws and ws.max_row and ws.max_column:
                                all_rows = []
                                for row in ws.iter_rows(values_only=True):
                                    all_rows.append([_format_cell_value(c) for c in row])
                                if all_rows:
                                    max_cols = max(len(r) for r in all_rows)
                                    for r in all_rows:
                                        while len(r) < max_cols:
                                            r.append('')
                                    header = all_rows[0]
                                    table_lines = ['| ' + ' | '.join(header) + ' |']
                                    table_lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
                                    for row in all_rows[1:]:
                                        table_lines.append('| ' + ' | '.join(row[:max_cols]) + ' |')
                                    # Map to first unmatched chart
                                    for cf in chart_files:
                                        cn = os.path.basename(cf)
                                        if cn not in chart_tables:
                                            chart_tables[cn] = '\n'.join(table_lines)
                                            break
                            wb.close()
                        except Exception:
                            continue

    except Exception:
        pass

    return chart_tables


def convert_docx_native(input_path, output_dir, config, args):
    """直接讀取 .docx，保留標題層級、表格、列表結構、圖片、圖表。"""
    try:
        import docx
    except ImportError:
        print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
        return False

    try:
        doc = docx.Document(str(input_path))
    except Exception as e:
        print(f"Error opening {input_path}: {e}", file=sys.stderr)
        return False

    print(f"Processing: {input_path}", file=sys.stderr)

    # Extract images and charts
    rid_to_md = _extract_docx_images(doc, input_path, output_dir, config)
    chart_tables = _extract_chart_data(input_path)

    # Build rId -> chart filename mapping from relationships
    rid_to_chart = {}
    for rel_id, rel in doc.part.rels.items():
        if "chart" in rel.reltype:
            try:
                chart_name = os.path.basename(rel.target_part.partname)
                rid_to_chart[rel_id] = chart_name
            except Exception:
                continue

    # XML namespaces for inline detection
    _A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    _R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    _C_NS = 'http://schemas.openxmlformats.org/drawingml/2006/chart'

    lines = []
    img_count = 0
    chart_count = 0

    # Build a flat iteration order: paragraphs and tables interleaved
    # by their position in document.body
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Check for inline images (a:blip)
            para_images = []
            for blip in element.iter(f'{{{_A_NS}}}blip'):
                rid = blip.get(f'{{{_R_NS}}}embed')
                if rid and rid in rid_to_md:
                    para_images.append(rid_to_md[rid])
                    img_count += 1

            # Check for inline charts (c:chart)
            para_charts = []
            for chart_ref in element.iter(f'{{{_C_NS}}}chart'):
                rid = chart_ref.get(f'{{{_R_NS}}}id')
                if rid and rid in rid_to_chart:
                    chart_name = rid_to_chart[rid]
                    if chart_name in chart_tables:
                        para_charts.append(chart_tables[chart_name])
                        chart_count += 1

            # Find matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()

                    if not text and not para_images and not para_charts:
                        lines.append('')
                        break

                    style_name = (para.style.name or '').lower()

                    # Output text
                    if text:
                        if style_name.startswith('heading'):
                            try:
                                level = int(style_name.replace('heading', '').strip())
                                level = min(max(level, 1), 6)
                            except ValueError:
                                level = 2
                            lines.append(f'{"#" * level} {text}')
                            lines.append('')
                        elif style_name.startswith('title'):
                            lines.append(f'# {text}')
                            lines.append('')
                        elif style_name.startswith('list'):
                            lines.append(f'- {text}')
                        else:
                            lines.append(text)
                            lines.append('')

                    # Insert image references
                    for img_md in para_images:
                        lines.append(img_md)
                        lines.append('')

                    # Insert chart tables
                    for chart_md in para_charts:
                        lines.append('')
                        lines.append(chart_md)
                        lines.append('')

                    break

        elif tag == 'tbl':
            # Find matching table object
            for table in doc.tables:
                if table._element is element:
                    rows_data = []
                    for row in table.rows:
                        row_cells = [cell.text.strip().replace('\n', ' ').replace('|', '\\|')
                                     for cell in row.cells]
                        rows_data.append(row_cells)

                    if rows_data and any(any(c for c in row) for row in rows_data):
                        max_cols = max(len(r) for r in rows_data)
                        for r in rows_data:
                            while len(r) < max_cols:
                                r.append('')

                        header = rows_data[0]
                        table_lines = ['| ' + ' | '.join(header) + ' |']
                        table_lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
                        for row in rows_data[1:]:
                            table_lines.append('| ' + ' | '.join(row[:max_cols]) + ' |')

                        lines.append('')
                        lines.append('\n'.join(table_lines))
                        lines.append('')
                    break

    if img_count:
        print(f"  Extracted {img_count} image(s)", file=sys.stderr)
    if chart_count:
        print(f"  Extracted {chart_count} chart(s) as table(s)", file=sys.stderr)

    full_md = '\n'.join(lines)

    # Clean up excessive blank lines
    full_md = re.sub(r'\n{3,}', '\n\n', full_md)
    if not full_md.endswith('\n'):
        full_md += '\n'

    # Frontmatter
    frontmatter = ""
    if not args.no_frontmatter and config.get("frontmatter", {}).get("include", True):
        fm_lines = ["---"]
        fm_lines.append(f"source_file: \"{Path(input_path).name}\"")
        fm_lines.append(f"converted_date: \"{datetime.now().strftime('%Y-%m-%d %H:%M')}\"")
        fm_lines.append(f"tool_version: \"{VERSION}\"")
        fm_lines.append("---")
        fm_lines.append("")
        frontmatter = '\n'.join(fm_lines)

    final_md = frontmatter + full_md

    # Fallback: if content is too sparse, retry via LibreOffice
    if not _check_content_sufficient(final_md):
        print(f"  Content too sparse ({Path(input_path).name}), falling back to LibreOffice...", file=sys.stderr)
        # Lazy import to avoid circular dependency
        from converters.libreoffice import convert_via_libreoffice
        return convert_via_libreoffice(input_path, output_dir, config, args)

    md_filename = Path(input_path).stem + '.md'
    md_path = output_dir / md_filename
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(final_md)

    print(f"  Output: {md_path}", file=sys.stderr)
    return True
